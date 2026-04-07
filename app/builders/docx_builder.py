"""DocxBuilder — builds the Executive Summary Word document.

Matches the handmade TPDDL MPG PM06 Executive Summary format:
  1. EXECUTIVE SUMMARY title
  2. CAPEX 25-26 - [work-type title]
  3. Applicant's Name / Address with Zone
  4. Indented field list (WBS, Order, CAPEX, Category, Sub-Head, Cost, etc.)
  5. Existing Scenario + (Checked by NEG)
  6. Proposed Scenario
  7. Major Material table (3 cols: Sr. No., Material, Quantity)
  8. BOQ and Cost Estimate boilerplate
  9. Cost table image from scheme copy
  10. Scope of the Scheme
  11. Remarks
  12. Timeline
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from app.builders.renderers import get_renderer
from app.builders.renderers.base_renderer import pick_pole_ref
from app.domain.enums import WorkType
from app.domain.models import Case
from app.infrastructure.formatting import format_indian_amount, get_capex_year

logger = logging.getLogger(__name__)


def _add_page_border(doc: Document) -> None:
    """Add double-line page border to all sections (matches handmade)."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
        sectPr.append(pgBorders)


def _set_highlight(run, color: str = 'yellow') -> None:
    """Set highlight colour on a run."""
    rPr = run._element.get_or_add_rPr()
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), color)
    rPr.append(hl)


def _set_cell_shading(cell, fill: str) -> None:
    """Set cell background fill colour (hex like 'C4BD97')."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: int = 11) -> None:
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text else "")
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Cambria"


def _short_capex_year(capex_year: str | None = None) -> str:
    """Return short CAPEX year like '25-26' instead of '2025-26'."""
    full = capex_year or get_capex_year()  # e.g. "2025-26"
    parts = full.split("-")
    return f"{parts[0][2:]}-{parts[1]}" if len(parts) == 2 and len(parts[0]) == 4 else full


def _space_material_desc(desc: str) -> str:
    """Insert spaces into SAP material descriptions for readability.

    'CABLE1.1KVAL4CX25SQMM' → 'CABLE 1.1KV AL 4CX25 SQMM'
    'TRANSFORMER250KVA3PH11KV/433VCU' → 'TRANSFORMER 250KVA 3PH 11KV/ 433V CU'
    'POLEPCC9MLONG160KG' → 'POLE PCC 9 M LONG 160 KG'
    """
    if not desc:
        return desc
    s = desc.strip()

    # Known SAP tokens — split these off as standalone words
    # Unit tokens (KV, KVA) should NOT be split from preceding digits (250KVA, 11KV)
    _UNIT_TOKENS = ["KVA", "KV"]  # order matters: KVA before KV
    _TOKENS = [
        "SQMM", "ARM", "XLPE", "PVC",
        "AL", "CU", "PH", "SWG", "GI", "ACB", "FDR",
        "PCC", "RCC", "LONG", "KG", "DIST",
        "POLYCARBONATE",
    ]
    # Unit tokens: only split from preceding LETTERS, not digits
    for tok in _UNIT_TOKENS:
        s = re.sub(rf"(?<=[A-Z])({tok})(?=[^A-Z]|$)", r" \1", s)
    # Other tokens: split from both letters and digits
    for tok in _TOKENS:
        s = re.sub(rf"(?<=[A-Z0-9/])({tok})(?=[^A-Z]|$)", r" \1", s)
        s = re.sub(rf"(?<=[a-z0-9])({tok})(?=[^A-Z]|$)", r" \1", s)

    # Insert space at letter→digit boundaries (but preserve "1.1", "4CX25" patterns)
    # Split: "CABLE1" → "CABLE 1", but keep "CX25" together (core-cross notation)
    s = re.sub(r"([A-Z]{2,}(?<!CX))(\d)", r"\1 \2", s)  # ≥2 letters (not CX) then digit
    # Don't split single letter + digit (like 4C, 3P) — that's "4-core", "3-phase"

    # Insert space at digit→letter for long letter runs (≥3 letters)
    # but NOT before KVA (keep 250KVA together) or CX
    s = re.sub(r"(\d)((?!KVA)[A-Z]{3,})", r"\1 \2", s)

    # Specific: digit followed by single M (meter) before space or end
    s = re.sub(r"(\d)(M)(?=\s|$)", r"\1 \2", s)

    # Insert space after / when followed by digit (11KV/433V → 11KV/ 433V)
    s = re.sub(r"/(\d)", r"/ \1", s)

    # Clean up multiple spaces
    s = re.sub(r" +", " ", s).strip()
    return s


def _select_key_materials(case: Case) -> list:
    """Select the key/major materials for the summary based on work type.

    - LT Standard: 1 primary cable
    - DT Augmentation: new transformer + ACBs (skip old/smaller transformer)
    - Others: top 1-2 by value
    """
    from app.domain.models import Material

    mats = case.materials or []
    if not mats:
        return []

    work_type = case.work_type or WorkType.LT_STANDARD

    if work_type == WorkType.DT_AUGMENTATION:
        # For DT Aug: pick the NEW (larger) transformer + major secondary material
        transformers = []
        acbs = []
        abc_cable = None
        for m in mats:
            desc = (m.description or "").upper()
            if "TRANSFORMER" in desc:
                match = re.search(r"(\d+)\s*KVA", desc)
                kva = int(match.group(1)) if match else 0
                transformers.append((kva, m))
            elif "ACB" in desc:
                acbs.append(m)
            elif "CABLE" in desc and "ABC" in desc:
                if abc_cable is None or float(m.quantity or 0) > float(abc_cable.quantity or 0):
                    abc_cable = m

        key = []
        # Pick the largest transformer (the new one)
        if transformers:
            transformers.sort(key=lambda t: t[0], reverse=True)
            key.append(transformers[0][1])
        # Prefer ABC cable as secondary when present with significant qty
        if abc_cable and float(abc_cable.quantity or 0) >= 100:
            key.append(abc_cable)
        else:
            key.extend(acbs[:1])
        return key if key else mats[:1]

    # LT Standard / ABC Wiring / others: pick primary cable + poles + dist box
    # For ABC_WIRING: prefer ABC cable over regular cable; include all pole types
    cables = []
    poles = []
    distbox = None
    for m in mats:
        desc = (m.description or "").upper()
        desc_compact = desc.replace(" ", "")
        if "CABLE" in desc:
            cables.append(m)
        elif desc.startswith("POLE") or "POLEPCC" in desc_compact or "POLERCC" in desc_compact:
            poles.append(m)
        elif ("DIST.BOX" in desc or "DIST BOX" in desc or "DISTBOX" in desc) and distbox is None:
            distbox = m

    # Pick primary cable: prefer ABC cable if present, else highest-quantity cable
    cable = None
    if cables:
        abc_cables = [c for c in cables if "ABC" in (c.description or "").upper()]
        if abc_cables:
            cable = abc_cables[0]
        else:
            # Pick highest quantity cable
            cables.sort(key=lambda c: float(c.quantity or 0), reverse=True)
            cable = cables[0]

    # Sort poles by quantity descending (most-used pole type first)
    poles.sort(key=lambda p: float(p.quantity or 0), reverse=True)

    # Order: cable first, then poles, then dist box (matches handmade layout)
    key = []
    if cable:
        key.append(cable)
    for p in poles:
        key.append(p)
    if distbox and (poles or work_type == WorkType.ABC_WIRING):
        key.append(distbox)
    return key if key else mats[:1]


def _build_scope_text(case: Case) -> str:
    """Build 'Scope of the Scheme' from key materials + tapping pole.

    LT Standard: '85 M CABLE 1.1KV AL 4CX25 SQMM ARM, are required at pole no. ...'
    DT Aug: '01no. TRANSFORMER 250KVA 3PH ... & 01no. LT ACB 400A ... is required at pole no. ...'
    """
    key_mats = _select_key_materials(case)

    if not key_mats and case.scope_of_work:
        return case.scope_of_work

    # Exclude dist box from scope text only for ABC_WIRING (handmade pattern)
    work_type = case.work_type or WorkType.LT_STANDARD
    if work_type == WorkType.ABC_WIRING:
        scope_mats = [
            m for m in key_mats
            if not any(kw in (m.description or "").upper() for kw in ("DIST.BOX", "DIST BOX", "DISTBOX"))
        ]
    else:
        scope_mats = key_mats
    pole = pick_pole_ref(case.tapping_pole, case.substation_name)
    pole_ref = f"pole no. {pole}"

    if work_type == WorkType.DT_AUGMENTATION:
        # DT Aug format: "01no. MATERIAL DESC & 01no. MATERIAL DESC is required at pole no. ..."
        parts = []
        for mat in scope_mats:
            desc = _space_material_desc((mat.description or "").upper())
            qty = mat.quantity
            qty_str = str(int(qty)) if qty and float(qty) == int(float(qty)) else "01"
            parts.append(f"{qty_str.zfill(2)}no. {desc}")
        mat_text = " & ".join(parts)
        return f"{mat_text} is required at {pole_ref}."
    else:
        # LT Standard format: "85 M CABLE ..., are required at pole no. ..."
        parts = []
        for mat in scope_mats:
            qty = mat.quantity
            unit = (mat.unit or "").upper()
            desc = _space_material_desc((mat.description or "").upper())
            if qty is not None:
                qty_str = str(int(qty)) if float(qty) == int(float(qty)) else str(qty)
                if unit in ("M", "MTR"):
                    parts.append(f"{qty_str} M {desc}")
                elif unit in ("NO", "NOS", "EA"):
                    parts.append(f"{qty_str.zfill(2)}NO {desc}")
                else:
                    parts.append(f"{qty_str} {unit} {desc}")
            else:
                parts.append(desc)

        if not parts:
            return case.scope_of_work or ""

        if len(parts) > 1:
            mat_text = ", ".join(parts[:-1]) + " and " + parts[-1]
        else:
            mat_text = parts[0] if parts else ""
        return f"{mat_text}, are required at {pole_ref} towards above mentioned address."


def _clean_applicant_name(raw: str | None) -> str:
    """Clean applicant name: remove 'Company' prefix, trailing periods."""
    if not raw:
        return ""
    s = str(raw).strip()
    # Remove "Company" prefix (e.g. "Company ADITYA SETH" → "ADITYA SETH")
    s = re.sub(r"^Company\s+", "", s, flags=re.I)
    # Remove trailing period
    s = s.rstrip(".")
    return s.strip()


def _clean_load(raw: str) -> str:
    """Clean load_applied value: strip category text and avoid kW duplication.

    '01KW E-DOM' → '1 kW'
    '60 KW'      → '60 kW'
    '5 KW E-NDLT' → '5 kW'
    '5KW E-VEHICLE' → '5 kW'
    '2 kw Domestic of both connection' → '2 kW'
    '01KVA E-DOMESTIC' → '1 kW'
    '01KW +01' → '2 kW'
    '18'         → '18 kW'
    """
    if not raw:
        return "N/A"
    s = str(raw).strip()

    # Handle multiple loads joined with + (e.g. "01KW +01 kW", "01KW +01")
    parts = re.split(r"\s*\+\s*", s)
    if len(parts) > 1:
        total = 0
        for part in parts:
            p = part.strip()
            p = re.sub(r"(?i)\s*k[wW]\b.*", "", p)
            p = re.sub(r"(?i)\s*KVA\b.*", "", p)
            p = p.strip().lstrip("0") or "0"
            try:
                total += float(p)
            except ValueError:
                pass
        if total > 0:
            s = str(int(total)) if total == int(total) else str(total)
            return f"{s} kW"

    # Remove supply-category suffixes: E-DOM, E-NDL, E-NDLT, E-DOMESTIC, E-NDS, E-COM, E-IND, E-AGR, E-VEHICLE, etc.
    s = re.sub(r"\s*E-[A-Z]+\b", "", s, flags=re.I)
    # Remove freeform text after the numeric+unit part ("Domestic of both connection", etc.)
    s = re.sub(r"(?i)\s*(domestic|connection|commercial|industrial|agricultural|vehicle).*", "", s)
    s = s.strip()
    # Strip any existing kW/KW/kw/KVA suffix to normalise
    s = re.sub(r"(?i)\s*k[wW]\s*$", "", s)
    s = re.sub(r"(?i)k[wW]$", "", s)          # handles "01KW" (no space)
    s = re.sub(r"(?i)\s*KVA\s*$", "", s)
    s = re.sub(r"(?i)KVA$", "", s)             # handles "01KVA"
    s = s.strip()
    # Strip leading zeros for clean display ("01" → "1", but keep "0.5" as-is)
    if s and s.isdigit() and len(s) > 1:
        s = s.lstrip("0") or "0"
    return f"{s} kW" if s else "N/A"


def _clean_address(raw: str) -> str:
    """Clean raw address from PM06 Excel (remove prefixes, mobile, etc.)."""
    if not raw:
        return raw
    s = raw
    # Remove "Supply Address:" prefix
    s = re.sub(r"(?i)^supply\s*address\s*:\s*", "", s)
    # Remove LANDMARK and everything after it (handles N/A, NA, ., comma, text)
    s = re.sub(r"(?i)\s*LANDMARK\s.*", "", s, flags=re.DOTALL)
    # Remove Mobile/MOB number and everything after
    s = re.sub(r"(?i)\s*(MOB\.?\s*NO\.?\s*[\d]+|Mobile\s*[-\u2013:].*).*", "", s, flags=re.DOTALL)
    # Remove Email and everything after
    s = re.sub(r"(?i)\s*(E-?mail\s*[-:].*).*", "", s, flags=re.DOTALL)
    # Remove Communication Address and everything after
    s = re.sub(r"(?i)\s*Communication\s*Address.*", "", s, flags=re.DOTALL)
    # Remove NEAR BY POLE references
    s = re.sub(r"(?i)\s*NEAR\s*BY\s*POLE\s*NO\.?\s*[\w\-/\.]*", "", s)
    s = re.sub(r"(?i)\s*NEAR\s*BY\s*POLE\s*[\w\-/\.]*", "", s)
    # Replace newlines with spaces
    s = s.replace("\n", " ").replace("\r", " ")
    # Remove leading dots/spaces
    s = re.sub(r"^[\s.]+", "", s)
    # Remove trailing commas/dots/spaces
    s = re.sub(r"[,\.\s]+$", "", s)
    # Collapse multiple spaces
    s = re.sub(r" +", " ", s).strip()
    return s


def _extract_area_name(address: str) -> str:
    """Extract area/village name from cleaned address for concise display.

    Looks for patterns like 'Village & Post Office Village Alipur'.
    Returns uppercase area name or empty string.
    """
    if not address:
        return ""
    addr_upper = address.upper()
    # Try "Village & Post Office [Village] X" pattern
    m = re.search(r"VILLAGE\s*&\s*POST\s+OFFICE\s+(?:VILLAGE\s+)?(\w+)", addr_upper)
    if m:
        return m.group(1).strip()
    # Try "VPO X" pattern
    m = re.search(r"VPO\s+(\w+)", addr_upper)
    if m:
        return m.group(1).strip()
    # Try last village mention: "Village X" not followed by &
    m = re.search(r"VILLAGE\s+(\w+)(?!\s*&)", addr_upper)
    if m and m.group(1) != "&":
        return m.group(1).strip()
    return ""


class DocxBuilder:
    """Builds an Executive Summary .docx from a Case object."""

    def build_summary(
        self,
        case: Case,
        output_path: str | Path,
        cost_table_image: Optional[bytes] = None,
    ) -> Path:
        """Generate the executive summary document.

        Returns the Path of the saved file.
        """
        output_path = Path(output_path)
        doc = Document()

        # -- Page margins (matches handmade) --
        for section in doc.sections:
            section.top_margin = 914400      # 1 inch
            section.bottom_margin = 742950   # ~0.585 inch
            section.left_margin = 914400     # ~0.72 inch
            section.right_margin = 571500    # ~0.45 inch

        # -- Double-line page border (matches handmade) --
        _add_page_border(doc)

        # -- Set default font --
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Cambria"
        font.size = Pt(11)

        # -- Tighten default spacing (matches handmade: no after-space, single line) --
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        work_type = case.work_type or WorkType.LT_STANDARD
        renderer = get_renderer(work_type)

        capex_short = _short_capex_year(case.capex_year)
        capex_full = case.capex_year or get_capex_year()
        notif = case.notification_no or "[Notification No]"

        # ── 1. EXECUTIVE SUMMARY title ──────────────────────────
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.line_spacing = 1.5
        run_title = p_title.add_run("EXECUTIVE SUMMARY")
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_title.font.name = "Cambria"

        # Spacers (match handmade: 3 blanks after title, first 2 with 1.5 line spacing)
        _b1 = doc.add_paragraph()
        _b1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _b1.paragraph_format.line_spacing = 1.5
        _b2 = doc.add_paragraph()
        _b2.paragraph_format.line_spacing = 1.5
        _b3 = doc.add_paragraph()
        _b3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── 2. CAPEX title line ─────────────────────────────────
        capex_title = renderer.capex_title(
            notification_no=notif,
            existing_dt_capacity=case.existing_dt_capacity,
            new_transformer_rating=case.new_transformer_rating,
            acb_description=case.acb_description,
        )
        p_capex = doc.add_paragraph()
        run_capex = p_capex.add_run(f"CAPEX {capex_short} - {capex_title}")
        run_capex.bold = True
        run_capex.font.size = Pt(11)
        run_capex.font.name = "Cambria"

        # ── 3. Applicant's Name ─────────────────────────────────
        applicant_name = _clean_applicant_name(case.applicant_name)
        p_name = doc.add_paragraph()
        rn_label = p_name.add_run("Applicant's Name: ")
        rn_label.bold = True
        rn_label.font.size = Pt(11)
        rn_label.font.name = "Cambria"
        rn_value = p_name.add_run(applicant_name or "N/A")
        rn_value.font.size = Pt(11)
        rn_value.font.name = "Cambria"

        # ── 4. Applicant's Address with Zone ────────────────────
        cleaned_addr = _clean_address(case.address or "")
        full_address = cleaned_addr
        # Notification no and zone go on separate line in the address block
        suffix_parts = []
        if case.notification_no:
            suffix_parts.append(f"vide Notification no. \u2013 {case.notification_no}.")
        zone_str = case.zone or ""
        district_str = case.district or ""
        if zone_str or district_str:
            zone_display = f"Zone {zone_str}"
            if district_str:
                zone_display += f" D-{district_str}"
            suffix_parts.append(zone_display)
        if suffix_parts:
            full_address = cleaned_addr + " " + " ".join(suffix_parts) if cleaned_addr else " ".join(suffix_parts)

        p_addr = doc.add_paragraph()
        ra_label = p_addr.add_run("Applicant's Address: ")
        ra_label.bold = True
        ra_label.font.size = Pt(11)
        ra_label.font.name = "Cambria"
        ra_value = p_addr.add_run(full_address)
        ra_value.font.size = Pt(11)
        ra_value.font.name = "Cambria"

        # Blank between address and field list
        doc.add_paragraph()

        # ── 5. Indented field list ──────────────────────────────
        fields = [
            ("WBS (Scheme) No.", case.wbs_no or "CE/N0000/XXXXX"),
            ("PM06 Order No", case.order_no or "N/A"),
            ("CAPEX", capex_full),
            ("Main Scheme Category", "Load Growth Schemes"),
            ("Sub-Head", renderer.sub_head()),
            ("Sd. Load Applied (kW)", _clean_load(case.load_applied) if case.load_applied else "N/A"),
            ("Estimated Cost( TPDDL)", f"\u20b9 {format_indian_amount(case.grand_total)}" if case.grand_total else "\u20b9 N/A"),
            ("R.R Charges", f"\u20b9 {format_indian_amount(case.rrc_total)}" if case.rrc_total else "\u20b9  XX,XX,XXX"),
            ("Budget Head", "TPDDL-NC"),
            ("Justification", "Existing/Proposed Scenario"),
        ]

        for label, value in fields:
            p_field = doc.add_paragraph(style="List Paragraph")
            # Set tab stops for clean colon/value alignment (proportional font)
            pPr = p_field._element.get_or_add_pPr()
            tabs_el = OxmlElement('w:tabs')
            for pos in ('4320', '4680'):
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'left')
                tab.set(qn('w:pos'), pos)
                tabs_el.append(tab)
            pPr.append(tabs_el)

            rl = p_field.add_run(label)
            rl.bold = True
            rl.font.size = Pt(11)
            rl.font.name = "Cambria"
            rt = p_field.add_run("\t:\t")
            rt.bold = True
            rt.font.size = Pt(11)
            rt.font.name = "Cambria"
            rv = p_field.add_run(str(value))
            rv.bold = True
            rv.font.size = Pt(11)
            rv.font.name = "Cambria"
            # Yellow highlight on WBS label+value (matches handmade)
            if label.startswith("WBS"):
                _set_highlight(rl, 'yellow')
                _set_highlight(rt, 'yellow')
                _set_highlight(rv, 'yellow')
            if label == "Budget Head":
                _set_highlight(rv, 'yellow')

        # Two blanks before Existing Scenario (match handmade, List Paragraph style)
        doc.add_paragraph(style="List Paragraph")
        _bs1 = doc.add_paragraph(style="List Paragraph")
        _bs1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── 6. Existing Scenario ────────────────────────────────
        p_es = doc.add_paragraph()
        p_es.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_es = p_es.add_run("Existing Scenario:")
        run_es.bold = True
        run_es.font.size = Pt(11)
        run_es.font.name = "Cambria"

        es_text = renderer.existing_scenario(
            notif,
            dt_loading=case.dt_loading,
            existing_dt_capacity=case.existing_dt_capacity,
            detailed_reason=case.detailed_reason,
        )
        # Clean any stray newlines and collapse whitespace
        es_text = re.sub(r'\s*\n\s*', ' ', es_text).strip()
        es_text = re.sub(r'  +', ' ', es_text)

        # Render as a single justified paragraph for clean formatting
        p_es_text = doc.add_paragraph()
        p_es_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_es_t = p_es_text.add_run(es_text)
        run_es_t.font.size = Pt(11)
        run_es_t.font.name = "Cambria"

        p_neg = doc.add_paragraph()
        p_neg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_neg = p_neg.add_run("(Checked by NEG)")
        run_neg.font.size = Pt(11)
        run_neg.font.name = "Cambria"
        _set_highlight(run_neg, 'yellow')

        # Spacer
        _bs2 = doc.add_paragraph()
        _bs2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── 7. Proposed Scenario ────────────────────────────────
        p_ps = doc.add_paragraph()
        p_ps.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_ps = p_ps.add_run("Proposed Scenario:")
        run_ps.bold = True
        run_ps.font.size = Pt(11)
        run_ps.font.name = "Cambria"

        # Blank between heading and proposed text (matches handmade)
        doc.add_paragraph()

        p_ps_text = doc.add_paragraph()
        run_ps_t = p_ps_text.add_run(
            renderer.proposed_scenario(
                tapping_pole=case.tapping_pole,
                existing_dt_capacity=case.existing_dt_capacity,
                new_transformer_rating=case.new_transformer_rating,
                acb_description=case.acb_description,
                substation_name=case.substation_name,
            )
        )
        run_ps_t.font.size = Pt(11)
        run_ps_t.font.name = "Cambria"

        # ── 8. Major Material table (3 cols) ────────────────────
        key_materials = _select_key_materials(case)
        if key_materials:
            _bs3 = doc.add_paragraph()
            _bs3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_mh = doc.add_paragraph()
            run_mh = p_mh.add_run("Major Material:")
            run_mh.bold = True
            run_mh.font.size = Pt(11)
            run_mh.font.name = "Cambria"

            table = doc.add_table(
                rows=1 + len(key_materials), cols=3, style="Table Grid"
            )
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set proportional column widths (matches handmade: ~13% / ~61% / ~26%)
            col_widths = [Inches(0.88), Inches(4.12), Inches(1.77)]  # ~63/297/128pt
            for row in table.rows:
                for ci, w in enumerate(col_widths):
                    row.cells[ci].width = w

            # Header — gold/tan fill matches handmade (C4BD97)
            for i, header in enumerate(["Sr. No.", "Material", "Quantity"]):
                _set_cell_text(table.rows[0].cells[i], header, bold=True)
                _set_cell_shading(table.rows[0].cells[i], "C4BD97")

            # Data rows
            for idx, mat in enumerate(key_materials):
                row = table.rows[idx + 1]
                _set_cell_text(row.cells[0], str(idx + 1))
                _set_cell_text(row.cells[1], _space_material_desc((mat.description or "").upper()))
                qty = mat.quantity
                qty_str = str(int(qty)) if qty and float(qty) == int(float(qty)) else str(qty or "")
                _set_cell_text(row.cells[2], qty_str)
                # Pink fill on first data row material cell — matches handmade
                if idx == 0:
                    _set_cell_shading(row.cells[1], "F2DCDB")

        # ── 9. BOQ and Cost Estimate boilerplate ────────────────
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        p_boq = doc.add_paragraph()
        run_boq_label = p_boq.add_run("BOQ and Cost Estimate:")
        run_boq_label.bold = True
        run_boq_label.font.size = Pt(9)
        run_boq_label.font.name = "Cambria"
        run_boq_text = p_boq.add_run(
            "The detailed quantities of various works to be executed under the scheme "
            "have been worked out based on a preliminary field survey and considerations "
            "discussed above. The detailed cost estimate is enclosed in the scheme."
        )
        run_boq_text.font.size = Pt(9)
        run_boq_text.font.name = "Cambria"

        # Spacers between BOQ and cost table (match handmade, 1.5 + JUSTIFY)
        for _ in range(6):
            _sp = doc.add_paragraph()
            _sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _sp.paragraph_format.line_spacing = 1.5

        # ── 10. Cost table image ────────────────────────────────
        p_cost_label = doc.add_paragraph()
        p_cost_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_cost_label.paragraph_format.line_spacing = 1.5
        run_cl = p_cost_label.add_run(
            "The summary of the cost estimate is given below :- "
        )
        run_cl.bold = True
        run_cl.font.size = Pt(11)
        run_cl.font.name = "Cambria"
        # Yellow highlight on "(from Scheme copy)" (matches handmade)
        run_cl2 = p_cost_label.add_run("(from Scheme copy)")
        run_cl2.bold = True
        run_cl2.font.size = Pt(11)
        run_cl2.font.name = "Cambria"
        _set_highlight(run_cl2, 'yellow')

        if cost_table_image:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_img = p_img.add_run()
            run_img.bold = True
            run_img.font.size = Pt(11)
            run_img.font.name = "Cambria"
            run_img.add_picture(io.BytesIO(cost_table_image), width=Inches(6.875))

        doc.add_paragraph()
        doc.add_paragraph()

        # ── 11. Scope of the Scheme ─────────────────────────────
        p_scope_heading = doc.add_paragraph()
        run_sh = p_scope_heading.add_run("Scope of the Scheme:")
        run_sh.bold = True
        run_sh.font.size = Pt(11)
        run_sh.font.name = "Cambria"

        scope_text = _build_scope_text(case)
        if scope_text:
            p_scope = doc.add_paragraph()
            run_scope = p_scope.add_run(scope_text)
            run_scope.font.size = Pt(11)
            run_scope.font.name = "Cambria"

        doc.add_paragraph(style="List Paragraph")

        # ── 12. Remarks ─────────────────────────────────────────
        p_remarks_heading = doc.add_paragraph()
        run_rh = p_remarks_heading.add_run("Remarks:")
        run_rh.bold = True
        run_rh.font.size = Pt(11)
        run_rh.font.name = "Cambria"

        doc.add_paragraph()

        p_remarks = doc.add_paragraph(style="List Paragraph")
        p_remarks.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_rm = p_remarks.add_run("RR and ROW  are required.")
        run_rm.font.size = Pt(11)
        run_rm.font.name = "Cambria"

        doc.add_paragraph()
        doc.add_paragraph(style="List Paragraph")

        # ── 13. Timeline ────────────────────────────────────────
        p_timeline = doc.add_paragraph(style="List Paragraph")
        # "7 days time" with cyan highlight
        run_tl1 = p_timeline.add_run("7 days time")
        run_tl1.font.size = Pt(11)
        run_tl1.font.name = "Cambria"
        _set_highlight(run_tl1, 'cyan')
        # Normal text
        run_tl2 = p_timeline.add_run(
            " is required for completion of the scheme ("
        )
        run_tl2.font.size = Pt(11)
        run_tl2.font.name = "Cambria"
        # "Subject to availability..." with yellow highlight + italic
        run_tl3 = p_timeline.add_run(
            "Subject to availability of necessary permission for Row "
            "from the agency concerned and availability of land; as applicable.)"
        )
        run_tl3.font.size = Pt(11)
        run_tl3.font.name = "Cambria"
        run_tl3.italic = True
        _set_highlight(run_tl3, 'yellow')

        # Trailing blanks (match handmade)
        _bs_end = doc.add_paragraph()
        _bs_end.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

        # ── Save ────────────────────────────────────────────────
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("Executive Summary saved \u2192 %s", output_path)
        return output_path